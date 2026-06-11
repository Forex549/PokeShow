"""
Genera heuristicas_pokeshow.docx con las heurísticas de los niveles 2, 4 y 5.
Las fórmulas se insertan como ecuaciones OMML nativas de Word.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── helpers OMML ─────────────────────────────────────────────────────────────

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def omml(tag, **attrs):
    el = OxmlElement(f"m:{tag}")
    for k, v in attrs.items():
        el.set(qn(f"m:{k}"), v)
    return el

def omml_r(text, italic=True, bold=False):
    """Crea un run matemático con texto."""
    run = omml("r")
    rpr = omml("rPr")
    sty = omml("sty")
    sty.set(qn("m:val"), "i" if italic and not bold else ("b" if bold else "p"))
    rpr.append(sty)
    run.append(rpr)
    t = omml("t")
    t.text = text
    run.append(t)
    return run

def omml_plain(text):
    return omml_r(text, italic=False)

def omml_italic(text):
    return omml_r(text, italic=True)

def build_omath(*children):
    """Envuelve en <m:oMath>."""
    math = OxmlElement("m:oMath")
    for c in children:
        math.append(c)
    return math

def make_para_with_equation(doc, math_el):
    """Inserta un párrafo que contiene solo la ecuación, centrado."""
    p = doc.add_paragraph()
    p.alignment = 1  # center
    pPr = p._p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    p._p.append(math_el)
    return p

def omml_fraction(num_children, den_children):
    """num / den como fracción."""
    f = omml("f")
    num = omml("num")
    for c in num_children:
        num.append(c)
    den = omml("den")
    for c in den_children:
        den.append(c)
    f.append(num)
    f.append(den)
    return f

def omml_sub(base_children, sub_children):
    """base_sub."""
    ss = omml("sSub")
    e = omml("e")
    for c in base_children:
        e.append(c)
    sub = omml("sub")
    for c in sub_children:
        sub.append(c)
    ss.append(e)
    ss.append(sub)
    return ss

def omml_sup(base_children, sup_children):
    """base^sup."""
    ss = omml("sSup")
    e = omml("e")
    for c in base_children:
        e.append(c)
    sup = omml("sup")
    for c in sup_children:
        sup.append(c)
    ss.append(e)
    ss.append(sup)
    return ss

def omml_func(name, arg_children):
    """func(arg)."""
    f = omml("func")
    fname = omml("fName")
    fname.append(omml_plain(name))
    e = omml("e")
    for c in arg_children:
        e.append(c)
    f.append(fname)
    f.append(e)
    return f

def omml_nary(char, sub_children, sup_children, e_children):
    """Símbolo ∑ / ∏ / max con límites."""
    nary = omml("nary")
    naryPr = omml("naryPr")
    chrEl = omml("chr")
    chrEl.set(qn("m:val"), char)
    naryPr.append(chrEl)
    nary.append(naryPr)
    sub = omml("sub")
    for c in sub_children:
        sub.append(c)
    sup = omml("sup")
    for c in sup_children:
        sup.append(c)
    e = omml("e")
    for c in e_children:
        e.append(c)
    nary.append(sub)
    nary.append(sup)
    nary.append(e)
    return nary

# ── estilos de texto ──────────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def heading3(doc, text):
    return doc.add_heading(text, level=3)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def add_table_penalties(doc):
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Estado"
    hdr[1].text = "Penalización"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rows = [
        ("slp (sueño)", "−1.0"),
        ("frz (congelado)", "−0.9"),
        ("par (parálisis)", "−0.5"),
        ("brn (quemadura)", "−0.4"),
        ("psn (veneno)", "−0.3"),
        ("confusion (volátil)", "−0.3"),
    ]
    for i, (estado, pen) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = estado
        table.rows[i].cells[1].text = pen
    doc.add_paragraph()

# ── ecuaciones ────────────────────────────────────────────────────────────────

def eq_score_matchup():
    """score = max_{m ∈ M} (power_m × eff_m × STAB_m)"""
    math = OxmlElement("m:oMath")

    # "score = "
    r_score = omml_italic("score")
    r_eq = omml_plain(" = ")
    math.append(r_score)
    math.append(r_eq)

    # max con subíndice m ∈ M
    nary = omml("nary")
    naryPr = omml("naryPr")
    chrEl = omml("chr"); chrEl.set(qn("m:val"), "max")
    limLocEl = omml("limLoc"); limLocEl.set(qn("m:val"), "subSup")
    naryPr.append(chrEl); naryPr.append(limLocEl)
    nary.append(naryPr)
    sub = omml("sub")
    sub.append(omml_italic("m")); sub.append(omml_plain(" ∈ ")); sub.append(omml_italic("M"))
    sup = omml("sup")  # vacío
    e = omml("e")

    # power_m × eff_m × STAB_m
    pow_sub = omml_sub([omml_italic("power")], [omml_italic("m")])
    eff_sub = omml_sub([omml_italic("eff")], [omml_italic("m")])
    stab_sub = omml_sub([omml_italic("STAB")], [omml_italic("m")])
    e.append(pow_sub)
    e.append(omml_plain(" × "))
    e.append(eff_sub)
    e.append(omml_plain(" × "))
    e.append(stab_sub)

    nary.append(sub); nary.append(sup); nary.append(e)
    math.append(nary)
    return math


def eq_stab():
    """STAB_m = 1.5 si tipo_m ∈ tipos(P), si no 1.0"""
    math = OxmlElement("m:oMath")
    stab = omml_sub([omml_italic("STAB")], [omml_italic("m")])
    math.append(stab)
    math.append(omml_plain(" = "))

    # Usar delimitador de llave para definición por casos
    d = omml("d")
    dPr = omml("dPr")
    begChr = omml("begChr"); begChr.set(qn("m:val"), "{")
    endChr = omml("endChr"); endChr.set(qn("m:val"), "")
    dPr.append(begChr); dPr.append(endChr)
    d.append(dPr)
    e = omml("e")
    e.append(omml_plain("1.5  si  tipo"))
    e.append(omml_sub([omml_italic("m")], []))
    e.append(omml_plain(" ∈ tipos("))
    e.append(omml_italic("P"))
    e.append(omml_plain("),   1.0  en otro caso"))
    d.append(e)
    math.append(d)
    return math


def eq_nivel4():
    """score = w1·f1 + w2·f2 + w3·f3 + w4·f4"""
    math = OxmlElement("m:oMath")
    math.append(omml_italic("score"))
    math.append(omml_plain(" = "))
    for i in range(1, 5):
        wi = omml_sub([omml_italic("w")], [omml_plain(str(i))])
        fi = omml_sub([omml_italic("f")], [omml_plain(str(i))])
        math.append(wi); math.append(omml_plain("·")); math.append(fi)
        if i < 4:
            math.append(omml_plain(" + "))
    return math


def eq_nivel5():
    """score = w1·f1 + w2·f2 + w3·f3 + w4·f4 + w5·f5"""
    math = OxmlElement("m:oMath")
    math.append(omml_italic("score"))
    math.append(omml_plain(" = "))
    for i in range(1, 6):
        wi = omml_sub([omml_italic("w")], [omml_plain(str(i))])
        fi = omml_sub([omml_italic("f")], [omml_plain(str(i))])
        math.append(wi); math.append(omml_plain("·")); math.append(fi)
        if i < 5:
            math.append(omml_plain(" + "))
    return math


def eq_f1():
    """f1 = (sum HP_i / sum maxHP_i)_IA - (sum HP_j / sum maxHP_j)_rival"""
    math = OxmlElement("m:oMath")
    math.append(omml_sub([omml_italic("f")], [omml_plain("1")]))
    math.append(omml_plain(" = "))

    def hp_ratio(label):
        # sum HP / sum maxHP con etiqueta de equipo
        num_sum = omml("nary")
        nPr = omml("naryPr")
        ch = omml("chr"); ch.set(qn("m:val"), "∑")
        nPr.append(ch); num_sum.append(nPr)
        num_sum.append(omml("sub")); num_sum.append(omml("sup"))
        e_num = omml("e"); e_num.append(omml_sub([omml_italic("HP")], [omml_italic("i")]))
        num_sum.append(e_num)

        den_sum = omml("nary")
        nPr2 = omml("naryPr")
        ch2 = omml("chr"); ch2.set(qn("m:val"), "∑")
        nPr2.append(ch2); den_sum.append(nPr2)
        den_sum.append(omml("sub")); den_sum.append(omml("sup"))
        e_den = omml("e")
        e_den.append(omml_sub([omml_italic("maxHP")], [omml_italic("i")]))
        den_sum.append(e_den)

        frac = omml_fraction([num_sum], [den_sum])
        sub_label = omml_sub([frac], [omml_plain(label)])
        return sub_label

    math.append(hp_ratio("IA"))
    math.append(omml_plain(" − "))
    math.append(hp_ratio("rival"))
    return math


def eq_f2():
    """f2 = vivos_IA/N_IA - vivos_rival/N_rival"""
    math = OxmlElement("m:oMath")
    math.append(omml_sub([omml_italic("f")], [omml_plain("2")]))
    math.append(omml_plain(" = "))
    math.append(omml_fraction(
        [omml_sub([omml_italic("vivos")], [omml_plain("IA")])],
        [omml_sub([omml_italic("N")], [omml_plain("IA")])]
    ))
    math.append(omml_plain(" − "))
    math.append(omml_fraction(
        [omml_sub([omml_italic("vivos")], [omml_plain("rival")])],
        [omml_sub([omml_italic("N")], [omml_plain("rival")])]
    ))
    return math


def eq_f3():
    """f3 = (spd_IA - spd_rival) / max(spd_IA, spd_rival)"""
    math = OxmlElement("m:oMath")
    math.append(omml_sub([omml_italic("f")], [omml_plain("3")]))
    math.append(omml_plain(" = "))
    num = [
        omml_sub([omml_italic("spd")], [omml_plain("IA")]),
        omml_plain(" − "),
        omml_sub([omml_italic("spd")], [omml_plain("rival")]),
    ]
    den = [
        omml_plain("max("),
        omml_sub([omml_italic("spd")], [omml_plain("IA")]),
        omml_plain(", "),
        omml_sub([omml_italic("spd")], [omml_plain("rival")]),
        omml_plain(")"),
    ]
    math.append(omml_fraction(num, den))
    return math


def eq_f4():
    """f4 = eff(mejor_IA → rival) - eff(mejor_rival → IA), normalizado [-1,1]"""
    math = OxmlElement("m:oMath")
    math.append(omml_sub([omml_italic("f")], [omml_plain("4")]))
    math.append(omml_plain(" = "))
    math.append(omml_plain("eff(mejor"))
    math.append(omml_sub([omml_plain("")], [omml_plain("IA")]))
    math.append(omml_plain(" → rival) − eff(mejor"))
    math.append(omml_sub([omml_plain("")], [omml_plain("rival")]))
    math.append(omml_plain(" → IA)"))
    return math


def eq_f5():
    """f5 = pen(rival) - pen(IA), donde pen = status_penalty + volatile_penalty"""
    math = OxmlElement("m:oMath")
    math.append(omml_sub([omml_italic("f")], [omml_plain("5")]))
    math.append(omml_plain(" = "))
    math.append(omml_plain("pen(rival) − pen(IA)"))
    return math


def eq_pen():
    """pen(P) = penalty_status(P) + penalty_volatile(P)"""
    math = OxmlElement("m:oMath")
    math.append(omml_plain("pen("))
    math.append(omml_italic("P"))
    math.append(omml_plain(") = "))
    math.append(omml_sub([omml_plain("penalty")], [omml_plain("status")]))
    math.append(omml_plain("("))
    math.append(omml_italic("P"))
    math.append(omml_plain(") + "))
    math.append(omml_sub([omml_plain("penalty")], [omml_plain("volatile")]))
    math.append(omml_plain("("))
    math.append(omml_italic("P"))
    math.append(omml_plain(")"))
    return math


def eq_fitness_n4():
    """fitness = (1·WR_rand + 2·WR_greedy + 3·WR_N3) / 6"""
    math = OxmlElement("m:oMath")
    math.append(omml_italic("fitness"))
    math.append(omml_plain(" = "))
    num = [
        omml_plain("1·"),
        omml_sub([omml_plain("WR")], [omml_plain("rand")]),
        omml_plain(" + 2·"),
        omml_sub([omml_plain("WR")], [omml_plain("greedy")]),
        omml_plain(" + 3·"),
        omml_sub([omml_plain("WR")], [omml_plain("N3")]),
    ]
    den = [omml_plain("6")]
    math.append(omml_fraction(num, den))
    return math


def eq_fitness_n5():
    """fitness = (1·WR_rand + 2·WR_greedy + 3·WR_N3 + 4·WR_N4) / 10"""
    math = OxmlElement("m:oMath")
    math.append(omml_italic("fitness"))
    math.append(omml_plain(" = "))
    num = [
        omml_plain("1·"),
        omml_sub([omml_plain("WR")], [omml_plain("rand")]),
        omml_plain(" + 2·"),
        omml_sub([omml_plain("WR")], [omml_plain("greedy")]),
        omml_plain(" + 3·"),
        omml_sub([omml_plain("WR")], [omml_plain("N3")]),
        omml_plain(" + 4·"),
        omml_sub([omml_plain("WR")], [omml_plain("N4")]),
    ]
    den = [omml_plain("10")]
    math.append(omml_fraction(num, den))
    return math


# ── documento ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Título principal
    title = doc.add_heading("Heurísticas de IA — PokeShow", 0)
    title.alignment = 1
    doc.add_paragraph(
        "Descripción formal de las heurísticas implementadas en los niveles 2, 4 y 5 "
        "del motor de decisión de la IA de batalla Pokémon."
    ).alignment = 1
    doc.add_paragraph()

    # ── NIVEL 2 ───────────────────────────────────────────────────────────────
    heading1(doc, "Nivel 2 — Mejor Opción (Greedy)")

    heading2(doc, "Propósito")
    body(doc,
        "El Nivel 2 implementa una estrategia puramente reactiva: evalúa cada movimiento "
        "disponible en el turno actual y elige el que produce el mayor daño inmediato "
        "contra el Pokémon rival activo. No existe lookahead ni modelado del oponente.")

    heading2(doc, "Función de evaluación de movimientos")
    body(doc,
        "Para cada movimiento m disponible se calcula el daño exacto mediante "
        "calculate_damage y se selecciona el máximo:")
    make_para_with_equation(doc, eq_score_matchup())
    body(doc, "donde:")
    bullet(doc, "powerₘ — potencia base del movimiento m")
    bullet(doc, "effₘ — multiplicador de efectividad de tipo del movimiento m contra el rival")
    bullet(doc, "STABₘ — bonificación por Same-Type Attack Bonus:")
    make_para_with_equation(doc, eq_stab())
    body(doc,
        "El resultado es determinista: no interviene aleatoriedad de críticos ni "
        "variación de daño.")

    heading2(doc, "Cambio forzado (Pokémon debilitado)")
    body(doc,
        "Cuando el Pokémon activo se debilita, la IA evalúa a cada sustituto candidato "
        "con la misma función de matchup anterior, aplicada sin calculate_damage para "
        "evitar efectos de aleatoriedad:")
    make_para_with_equation(doc, eq_score_matchup())
    body(doc, "Se envía al candidato con mayor score contra el Pokémon rival activo.")

    heading2(doc, "Limitaciones")
    bullet(doc, "No anticipa movimientos del rival (sin lookahead).")
    bullet(doc, "Ignora velocidad, supervivencia del equipo y estados alterados.")
    bullet(doc, "No detecta situaciones donde conservar HP es superior a atacar.")
    doc.add_paragraph()

    # ── NIVEL 4 ───────────────────────────────────────────────────────────────
    heading1(doc, "Nivel 4 — Minimax con 4 Variables (pesos GA)")

    heading2(doc, "Propósito")
    body(doc,
        "El Nivel 4 combina Minimax con poda Alfa-Beta (profundidad 3) y una heurística "
        "de cuatro factores cuyos pesos son optimizados por un Algoritmo Genético. "
        "El cromosoma de cada individuo es:")
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("[w₁, w₂, w₃, w₄]").bold = True

    heading2(doc, "Factores de la heurística")
    body(doc,
        "Cada factor está normalizado al rango [−1, 1]. Un valor positivo indica ventaja "
        "para la IA; negativo indica ventaja para el rival.")

    heading3(doc, "f₁ — Ratio de HP")
    body(doc, "Diferencia de porcentaje de vida total entre equipos:")
    make_para_with_equation(doc, eq_f1())

    heading3(doc, "f₂ — Pokémon vivos")
    body(doc, "Diferencia relativa de Pokémon con HP > 0 en cada equipo:")
    make_para_with_equation(doc, eq_f2())

    heading3(doc, "f₃ — Velocidad")
    body(doc, "Diferencia relativa de velocidad entre los Pokémon activos:")
    make_para_with_equation(doc, eq_f3())

    heading3(doc, "f₄ — Ventaja de tipo")
    body(doc,
        "Diferencia entre la efectividad del mejor movimiento de la IA sobre el rival y "
        "la del mejor movimiento del rival sobre la IA, normalizada a [−1, 1]:")
    make_para_with_equation(doc, eq_f4())

    heading2(doc, "Puntuación final")
    make_para_with_equation(doc, eq_nivel4())

    heading2(doc, "Función objetivo (fitness)")
    body(doc,
        "El GA optimiza la tasa de victorias ponderada contra tres rivales de distinta "
        "dificultad. Ganarle al rival más difícil aporta más al fitness:")
    make_para_with_equation(doc, eq_fitness_n4())
    body(doc, "donde WRₓ es la win-rate contra el rival x en n_battles batallas.")

    heading2(doc, "Cambio forzado (Pokémon debilitado)")
    body(doc,
        "Al debilitarse el Pokémon activo, la IA simula el estado del equipo con cada "
        "sustituto disponible y aplica la heurística completa (f₁…f₄) con los pesos "
        "actuales. Se selecciona el candidato con mayor puntuación.")

    heading2(doc, "Parámetros del GA")
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    params = [
        ("Parámetro", "Valor"),
        ("Población", "40 individuos"),
        ("Generaciones", "20"),
        ("Batallas por individuo", "10 × 3 rivales"),
        ("Selección", "Torneo (k = 3)"),
        ("Cruce", "Uniforme (p = 0.8)"),
        ("Mutación", "Gaussiana (σ = 0.5, p = 0.2 por gen)"),
        ("Elitismo", "Top 2 pasan directamente"),
    ]
    for i, (k, v) in enumerate(params):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        if i == 0:
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

    heading2(doc, "Limitaciones")
    bullet(doc, "Los pesos son estáticos durante el combate.")
    bullet(doc, "No considera estados alterados (parálisis, sueño, etc.).")
    bullet(doc, "No evalúa el cambio voluntario de Pokémon.")
    bullet(doc, "El GA ajusta importancias, no descubre nuevos criterios.")
    doc.add_paragraph()

    # ── NIVEL 5 ───────────────────────────────────────────────────────────────
    heading1(doc, "Nivel 5 — Minimax con 5 Variables + Cambio Voluntario (pesos GA)")

    heading2(doc, "Propósito")
    body(doc,
        "El Nivel 5 extiende el Nivel 4 con dos capacidades adicionales: un quinto "
        "factor que penaliza estados alterados y la evaluación de cambios voluntarios "
        "de Pokémon dentro del árbol de búsqueda Minimax. El cromosoma es:")
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("[w₁, w₂, w₃, w₄, w₅]").bold = True

    heading2(doc, "Factores de la heurística")
    body(doc, "Los cuatro primeros factores (f₁…f₄) son idénticos al Nivel 4. Se agrega:")

    heading3(doc, "f₅ — Estados alterados")
    body(doc,
        "Diferencia de penalización por estado entre el Pokémon rival y el Pokémon "
        "propio. Un valor positivo indica que el rival sufre más restricciones:")
    make_para_with_equation(doc, eq_f5())
    body(doc, "donde:")
    make_para_with_equation(doc, eq_pen())
    body(doc, "Penalizaciones aplicadas:")
    add_table_penalties(doc)

    heading2(doc, "Puntuación final")
    make_para_with_equation(doc, eq_nivel5())

    heading2(doc, "Cambio voluntario en el árbol Minimax")
    body(doc,
        "En cada nodo del árbol, la IA evalúa tanto movimientos de ataque como ramas "
        "de cambio voluntario: simula enviar a cada sustituto disponible y aplica la "
        "heurística completa (f₁…f₅). El Minimax retorna la acción de mayor valor "
        "esperado, que puede ser un movimiento o un índice de Pokémon.")
    body(doc, "Esto habilita comportamientos como:")
    bullet(doc, "Retirar un Pokémon debilitado antes de que muera.")
    bullet(doc, "Cambiar a quien tenga mejor matchup de tipos.")
    bullet(doc, "Explotar un rival con estado alterado enviando al Pokémon más efectivo.")

    heading2(doc, "Función objetivo (fitness)")
    body(doc,
        "El GA optimiza la tasa de victorias ponderada contra cuatro rivales. "
        "Ganarle al Nivel 4 aporta cuatro veces más que ganarle al aleatorio:")
    make_para_with_equation(doc, eq_fitness_n5())

    heading2(doc, "Parámetros del GA")
    table2 = doc.add_table(rows=8, cols=2)
    table2.style = "Table Grid"
    params2 = [
        ("Parámetro", "Valor"),
        ("Población", "40 individuos"),
        ("Generaciones", "20"),
        ("Batallas por individuo", "10 × 4 rivales"),
        ("Selección", "Torneo (k = 3)"),
        ("Cruce", "Uniforme gen a gen (p = 0.8)"),
        ("Mutación", "Gaussiana (σ = 0.5, p = 0.2 por gen)"),
        ("Elitismo", "Top 2 pasan directamente"),
    ]
    for i, (k, v) in enumerate(params2):
        row = table2.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        if i == 0:
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

    heading2(doc, "Aprendizaje implícito de la política de cambios")
    body(doc,
        "El GA no aprende explícitamente cuándo cambiar de Pokémon. Lo hace de forma "
        "indirecta: los pesos [w₁…w₅] definen cuán atractivo es un estado post-cambio "
        "versus un estado post-ataque dentro del árbol Minimax. La presión evolutiva "
        "ajusta implícitamente la agresividad de la política de cambios.")

    heading2(doc, "Limitaciones")
    bullet(doc, "La profundidad del Minimax es fija (3 plies); no se adapta al tiempo.")
    bullet(doc, "Los pesos son estáticos durante el combate.")
    bullet(doc, "El rival modelado en el árbol solo ataca; no realiza cambios voluntarios.")
    bullet(doc, "El GA ajusta importancias, no descubre nuevos criterios heurísticos.")

    out = "heuristicas_pokeshow.docx"
    doc.save(out)
    print(f"Documento generado: {out}")

if __name__ == "__main__":
    build()
